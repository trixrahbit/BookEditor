"""
DOCX Importer - Convert Word documents into novel structure
"""

from typing import List, Dict, Tuple
from pathlib import Path
import re

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from models.project import Scene, Chapter, Part
from db_manager import DatabaseManager


class DocxImporter:
    """Import and parse DOCX files into novel structure"""

    def __init__(self):
        self.chapters = []
        self.scenes = []

    def import_docx(self, file_path: str, db_manager: DatabaseManager, project_id: str) -> Tuple[int, int, int]:
        """
        Import a DOCX file and convert it into novel structure

        Returns: (parts_created, chapters_created, scenes_created)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

        doc = Document(file_path)

        # Parse document structure with proper formatting
        chapters_data = self._parse_document(doc)

        # Create database entries
        parts_count = 0
        chapters_count = 0
        scenes_count = 0

        for chapter_data in chapters_data:
            # Create chapter
            chapter = Chapter(
                name=chapter_data['title'],
                summary='',
                parent_id=None,
                order=chapters_count
            )
            db_manager.save_item(project_id, chapter)
            chapters_count += 1

            # Create scenes for this chapter
            for scene_index, scene_data in enumerate(chapter_data['scenes']):
                scene = Scene(
                    name=scene_data['title'],
                    content=scene_data['content'],
                    parent_id=chapter.id,
                    order=scene_index,
                    word_count=self._count_words_in_html(scene_data['content'])
                )
                db_manager.save_item(project_id, scene)
                scenes_count += 1

        return parts_count, chapters_count, scenes_count

    def _count_words_in_html(self, html: str) -> int:
        """Count words in HTML content"""
        # Strip HTML tags
        text = re.sub(r'<[^>]+>', ' ', html)
        words = [w for w in text.split() if w]
        return len(words)

    def _parse_document(self, doc: Document) -> List[Dict]:
        """
        Parse document into chapters with scenes

        Rules:
        - Heading 1 = Chapter
        - Heading 2/3 = Scene within chapter
        - Scene breaks (***) = Split into multiple scenes
        - Default = One scene per chapter
        """
        chapters = []
        current_chapter = None
        current_scene_title = None
        current_paragraphs = []
        skip_first_heading = True

        for para in doc.paragraphs:
            style = para.style.name
            text = para.text.strip()

            if style == 'Heading 1':
                # Skip first heading (book title)
                if skip_first_heading:
                    skip_first_heading = False
                    continue

                # Save previous scene and chapter
                if current_chapter:
                    self._save_scene(current_chapter, current_scene_title, current_paragraphs)
                    chapters.append(current_chapter)

                # Start new chapter
                current_chapter = {
                    'title': text,
                    'scenes': []
                }
                current_scene_title = None
                current_paragraphs = []

            elif style in ['Heading 2', 'Heading 3']:
                # Heading 2/3 = explicit scene marker

                # Create chapter if needed
                if not current_chapter:
                    current_chapter = {
                        'title': 'Chapter 1',
                        'scenes': []
                    }

                # Save previous scene
                if current_scene_title or current_paragraphs:
                    self._save_scene(current_chapter, current_scene_title, current_paragraphs)

                # Start new scene with this heading as title
                current_scene_title = text
                current_paragraphs = []

            elif text:
                # Regular paragraph

                # Create chapter if needed
                if not current_chapter:
                    current_chapter = {
                        'title': 'Chapter 1',
                        'scenes': []
                    }

                current_paragraphs.append(para)

        # Save final scene and chapter
        if current_chapter:
            self._save_scene(current_chapter, current_scene_title, current_paragraphs)
            chapters.append(current_chapter)

        # Process scene breaks
        return self._process_scene_breaks(chapters)

    def _save_scene(self, chapter: Dict, scene_title: str, paragraphs: List):
        """Save accumulated paragraphs as a scene"""
        if not paragraphs:
            return

        content_html = self._paragraphs_to_html(paragraphs)

        # Generate scene title if not provided
        if not scene_title:
            scene_number = len(chapter['scenes']) + 1
            scene_title = f"{chapter['title']} - Scene {scene_number}"

        chapter['scenes'].append({
            'title': scene_title,
            'content': content_html
        })

        paragraphs.clear()

    def _paragraphs_to_html(self, paragraphs: List) -> str:
        """Convert paragraphs to HTML preserving formatting and spacing"""
        html_parts = []

        for para in paragraphs:
            para_html = '<p>'

            for run in para.runs:
                text = run.text
                if not text:
                    continue

                # Escape HTML special characters but preserve spaces
                text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                # Preserve formatting
                if run.bold and run.italic:
                    text = f'<strong><em>{text}</em></strong>'
                elif run.bold:
                    text = f'<strong>{text}</strong>'
                elif run.italic:
                    text = f'<em>{text}</em>'
                elif run.underline:
                    text = f'<u>{text}</u>'

                para_html += text

            para_html += '</p>'
            html_parts.append(para_html)

        return '\n'.join(html_parts)

    def _process_scene_breaks(self, chapters: List[Dict]) -> List[Dict]:
        """
        Process scene break markers (***) to split scenes
        """
        for chapter in chapters:
            new_scenes = []

            for scene in chapter['scenes']:
                content = scene['content']

                # Check for scene break markers
                split_scenes = self._split_by_scene_breaks(content)

                if len(split_scenes) > 1:
                    # Multiple scenes found - number them
                    base_title = scene['title']
                    # Remove existing " - Scene X" to avoid double numbering
                    if ' - Scene ' in base_title:
                        base_title = base_title.rsplit(' - Scene ', 1)[0]

                    for i, scene_content in enumerate(split_scenes, 1):
                        new_scenes.append({
                            'title': f"{base_title} - Scene {i}",
                            'content': scene_content
                        })
                else:
                    # Single scene - keep as is
                    new_scenes.append(scene)

            chapter['scenes'] = new_scenes

        return chapters

    def _split_by_scene_breaks(self, content: str) -> List[str]:
        """Split HTML content by scene break markers"""
        patterns = [
            r'<p>\s*\*\s*\*\s*\*\s*</p>',
            r'<p>\s*-\s*-\s*-\s*</p>',
            r'<p>\s*#\s*#\s*#\s*</p>',
        ]

        for pattern in patterns:
            parts = re.split(pattern, content)
            if len(parts) > 1:
                return [part.strip() for part in parts if part.strip()]

        # No breaks found
        return [content]

    def estimate_structure(self, file_path: str) -> Dict[str, int]:
        """
        Estimate what will be created without actually importing
        """
        if not DOCX_AVAILABLE:
            return {'error': 'python-docx not installed'}

        doc = Document(file_path)
        chapters = self._parse_document(doc)

        total_scenes = sum(len(ch['scenes']) for ch in chapters)
        total_words = 0

        for chapter in chapters:
            for scene in chapter['scenes']:
                total_words += self._count_words_in_html(scene['content'])

        counts = {
            'parts': 0,
            'chapters': len(chapters),
            'scenes': total_scenes,
            'total_words': total_words
        }

        return counts


class ImportDialog:
    """Helper for import UI dialogs"""

    @staticmethod
    def show_import_preview(file_path: str) -> Dict[str, int]:
        """Show preview of what will be imported"""
        importer = DocxImporter()
        return importer.estimate_structure(file_path)

    @staticmethod
    def perform_import(file_path: str, db_manager: DatabaseManager,
                      project_id: str) -> Tuple[int, int, int]:
        """Perform the actual import"""
        importer = DocxImporter()
        return importer.import_docx(file_path, db_manager, project_id)