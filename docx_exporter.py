"""
DOCX Exporter - Convert novel structure into Word documents

Rules:
- Heading 1 = Chapter
- Heading 3 = Scene
- Scene content is HTML (subset) and is converted to docx paragraphs + runs
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import List, Optional, Tuple
from pathlib import Path
import re
import html as html_lib

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from models.project import Scene, Chapter, Part  # Part optional / unused unless you want it
from db_manager import DatabaseManager


# -----------------------------
# HTML -> DOCX conversion utils
# -----------------------------

_TAG_TOKEN_RE = re.compile(r"(<[^>]+>)")
_TAG_RE = re.compile(r"^<\s*/?\s*([a-zA-Z0-9]+)")

def _strip_tags(s: str) -> str:
    return re.sub(r"<[^>]+>", "", s)

def _normalize_whitespace(s: str) -> str:
    # Keep intentional spacing reasonable; docx collapses multiple spaces visually anyway.
    return s.replace("\r\n", "\n").replace("\r", "\n")

def _html_to_paragraph_text_blocks(html: str) -> List[str]:
    """
    Split HTML into paragraph blocks.
    Supports:
      - <p> ... </p>
      - If no <p>, treat as one paragraph
    """
    html = _normalize_whitespace(html or "").strip()
    if not html:
        return []

    # Grab content inside <p> tags
    ps = re.findall(r"<p[^>]*>(.*?)</p>", html, flags=re.IGNORECASE | re.DOTALL)
    if ps:
        return [p.strip() for p in ps if p.strip()]

    # Fallback: treat whole as one paragraph
    return [html]


@dataclass
class RunStyle:
    bold: bool = False
    italic: bool = False
    underline: bool = False


def _apply_run_style(run, style: RunStyle):
    run.bold = style.bold
    run.italic = style.italic
    run.underline = style.underline


def _add_styled_runs_from_html(paragraph, html_fragment: str):
    """
    Convert a subset of inline HTML into docx runs on a single paragraph.
    Supported inline tags:
      <strong>/<b>, <em>/<i>, <u>, <br>
    Anything else is stripped but text is preserved.
    """
    frag = html_fragment or ""
    frag = frag.replace("&nbsp;", " ")
    frag = html_lib.unescape(frag)

    tokens = _TAG_TOKEN_RE.split(frag)

    style_stack: List[RunStyle] = [RunStyle()]
    current = style_stack[-1]

    def push(**kwargs):
        nonlocal current
        new_style = RunStyle(
            bold=kwargs.get("bold", current.bold),
            italic=kwargs.get("italic", current.italic),
            underline=kwargs.get("underline", current.underline),
        )
        style_stack.append(new_style)
        current = new_style

    def pop():
        nonlocal current
        if len(style_stack) > 1:
            style_stack.pop()
        current = style_stack[-1]

    for tok in tokens:
        if not tok:
            continue

        if tok.startswith("<") and tok.endswith(">"):
            tag_match = _TAG_RE.match(tok.strip())
            if not tag_match:
                continue

            tag = tag_match.group(1).lower()
            is_close = tok.strip().startswith("</")

            if tag in ("br",):
                # Line break inside a paragraph
                paragraph.add_run("\n")
                continue

            if is_close:
                if tag in ("strong", "b", "em", "i", "u"):
                    pop()
                continue

            # opening tags
            if tag in ("strong", "b"):
                push(bold=True)
            elif tag in ("em", "i"):
                push(italic=True)
            elif tag == "u":
                push(underline=True)
            else:
                # Unknown tag: ignore
                continue
        else:
            # Plain text
            text = tok
            if text:
                run = paragraph.add_run(text)
                _apply_run_style(run, current)


# -----------------------------
# Exporter
# -----------------------------

class DocxExporter:
    """Export a project's chapter/scene structure to a DOCX document."""

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

    def export_project(
        self,
        db_manager: DatabaseManager,
        project_id: str,
        output_path: str,
        book_title: Optional[str] = None,
        include_scene_breaks: bool = False,
    ) -> Tuple[int, int]:
        """
        Export the whole project into a DOCX file.

        Returns: (chapters_exported, scenes_exported)
        """
        doc = Document()

        # Optional title page-ish top heading (not Heading 1 to avoid importer skipping logic)
        if book_title:
            title_para = doc.add_paragraph(book_title)
            title_para.style = "Title"
            doc.add_paragraph("")  # spacer

        chapters = self._load_chapters(db_manager, project_id)

        chapters_exported = 0
        scenes_exported = 0

        for ch in chapters:
            chapters_exported += 1
            ch_title = getattr(ch, "name", None) or f"Chapter {chapters_exported}"

            # Chapter heading
            para = doc.add_paragraph(ch_title)
            para.style = "Heading 1"

            # Pull scenes for chapter
            scenes = self._load_scenes_for_chapter(db_manager, project_id, ch.id)

            if not scenes:
                # Keep structure even if empty
                doc.add_paragraph("")
                continue

            for s_idx, sc in enumerate(scenes, start=1):
                scenes_exported += 1
                scene_title = getattr(sc, "name", None) or f"{ch_title} - Scene {s_idx}"

                # Scene heading
                spara = doc.add_paragraph(scene_title)
                spara.style = "Heading 3"

                # Scene content (HTML -> docx)
                content_html = getattr(sc, "content", "") or ""
                self._write_scene_content(doc, content_html)

                if include_scene_breaks and s_idx < len(scenes):
                    br = doc.add_paragraph("***")
                    br.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("")  # spacer between chapters

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        return chapters_exported, scenes_exported

    # -------- data loading helpers --------

    def _load_chapters(self, db_manager: DatabaseManager, project_id: str) -> List[Chapter]:
        """
        Load chapters in display order. Supports multiple DBManager APIs.
        """

        # 1) Preferred: children query by type
        if hasattr(db_manager, "get_children"):
            chapters = db_manager.get_children(project_id, parent_id=None, item_type="chapter")
            return sorted(chapters, key=lambda c: getattr(c, "order", 0))

        # 2) Alternate: get by type method
        if hasattr(db_manager, "get_items_by_type"):
            chapters = db_manager.get_items_by_type(project_id, "chapter", parent_id=None)
            return sorted(chapters, key=lambda c: getattr(c, "order", 0))

        # 3) Fallback: pull all items and filter
        all_items = self._load_all_items(db_manager, project_id)
        chapters = [i for i in all_items if self._is_type(i, "chapter") and getattr(i, "parent_id", None) in (None, "")]
        return sorted(chapters, key=lambda c: getattr(c, "order", 0))

    def _load_scenes_for_chapter(self, db_manager: DatabaseManager, project_id: str, chapter_id: str) -> List[Scene]:
        """
        Load scenes for a chapter in display order. Supports multiple DBManager APIs.
        """

        # 1) Preferred: children query by type
        if hasattr(db_manager, "get_children"):
            scenes = db_manager.get_children(project_id, parent_id=chapter_id, item_type="scene")
            return sorted(scenes, key=lambda s: getattr(s, "order", 0))

        # 2) Alternate: get by type method
        if hasattr(db_manager, "get_items_by_type"):
            scenes = db_manager.get_items_by_type(project_id, "scene", parent_id=chapter_id)
            return sorted(scenes, key=lambda s: getattr(s, "order", 0))

        # 3) Fallback: pull all items and filter
        all_items = self._load_all_items(db_manager, project_id)
        scenes = [i for i in all_items if self._is_type(i, "scene") and getattr(i, "parent_id", None) == chapter_id]
        return sorted(scenes, key=lambda s: getattr(s, "order", 0))

    def _load_all_items(self, db_manager: DatabaseManager, project_id: str):
        """
        Attempt to load all items for the project using various DBManager APIs.
        """
        if hasattr(db_manager, "get_items"):
            return db_manager.get_items(project_id)

        if hasattr(db_manager, "load_items"):
            return db_manager.load_items(project_id)

        if hasattr(db_manager, "get_project_items"):
            return db_manager.get_project_items(project_id)

        if hasattr(db_manager, "list_project_items"):
            return db_manager.list_project_items(project_id)

        raise AttributeError(
            "DatabaseManager must provide one of: get_children(), get_items_by_type(), "
            "or a bulk loader like get_items()/load_items()/get_project_items()."
        )

    def _is_type(self, item, type_name: str) -> bool:
        """
        Robustly detect item type across different models.
        """
        # Common patterns:
        # - item.type == "chapter"/"scene"
        # - item.item_type == "chapter"/"scene"
        # - class name Chapter/Scene
        t = getattr(item, "type", None) or getattr(item, "item_type", None)
        if isinstance(t, str) and t.lower() == type_name.lower():
            return True
        return item.__class__.__name__.lower() == type_name.lower()


    # -------- writing helpers --------

    def _write_scene_content(self, doc: Document, content_html: str):
        """
        Convert stored HTML content into docx paragraphs and runs.
        Supports paragraphs (<p>) and inline styles (<strong>, <em>, <u>, <br>).
        """
        blocks = _html_to_paragraph_text_blocks(content_html)

        if not blocks:
            doc.add_paragraph("")  # keep spacing consistent
            return

        for block in blocks:
            # Some stored HTML might already contain nested <p> or weirdness. Strip outer <p> if present.
            block = block.strip()
            if not block:
                doc.add_paragraph("")
                continue

            # Create a new paragraph and add styled runs
            p = doc.add_paragraph()
            _add_styled_runs_from_html(p, block)


# -----------------------------
# Convenience helper
# -----------------------------

class ExportDialog:
    """Helper for export UI dialogs"""

    @staticmethod
    def perform_export(
        db_manager: DatabaseManager,
        project_id: str,
        output_path: str,
        book_title: Optional[str] = None,
        include_scene_breaks: bool = False,
    ) -> Tuple[int, int]:
        exporter = DocxExporter()
        return exporter.export_project(
            db_manager=db_manager,
            project_id=project_id,
            output_path=output_path,
            book_title=book_title,
            include_scene_breaks=include_scene_breaks,
        )
